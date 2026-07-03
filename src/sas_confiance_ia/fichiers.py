"""Extraction de texte des fichiers déposés (Lot 15).

.txt / .md / .csv en texte brut (UTF-8, repli latin-1), .docx via
python-docx (paragraphes ET tableaux : c'est dans les tableaux que vivent
les NIR), .pdf textuel via pypdf. Un PDF sans texte extractible (scanné)
est refusé explicitement : pas d'OCR en v1, un refus documenté vaut mieux
qu'un texte vide qui laisserait croire à une pseudonymisation réussie.

Le PDF caviardé (PyMuPDF, licence AGPL contaminante) est exclu de la v1 :
arbitrage Q6, docs/specs/QUESTIONS.md.
"""

import io
from collections import Counter
from collections.abc import Callable
from pathlib import PurePosixPath

EXTENSIONS_TEXTE = {".txt", ".md", ".csv"}
EXTENSIONS_SUPPORTEES = EXTENSIONS_TEXTE | {".docx", ".pdf"}


class ErreurFichier(ValueError):
    """Fichier refusé : format, lisibilité ou contenu inexploitable."""


class FormatNonSupporte(ErreurFichier):
    pass


class FichierIllisible(ErreurFichier):
    pass


class PdfSansTexte(ErreurFichier):
    pass


def _texte_brut(octets: bytes) -> str:
    try:
        return octets.decode("utf-8")
    except UnicodeDecodeError:
        # Repli documenté : les exports Windows arrivent encore en latin-1.
        return octets.decode("latin-1")


def _ouvrir_docx(octets: bytes):
    import zipfile

    from docx import Document
    from docx.opc.exceptions import OpcError

    try:
        return Document(io.BytesIO(octets))
    except (OpcError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise FichierIllisible(f"fichier .docx illisible ({type(exc).__name__})") from exc


def _texte_docx(octets: bytes) -> str:
    document = _ouvrir_docx(octets)
    morceaux = [p.text for p in document.paragraphs]
    # Les tableaux après les paragraphes : l'entrelacement d'origine n'est
    # pas préservé (limite documentée), aucune cellule n'est perdue.
    for table in document.tables:
        for ligne in table.rows:
            morceaux.append("\t".join(cellule.text for cellule in ligne.cells))
    return "\n".join(morceaux)


def _texte_pdf(octets: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        lecteur = PdfReader(io.BytesIO(octets))
        pages = [page.extract_text() or "" for page in lecteur.pages]
    except (PyPdfError, ValueError, KeyError) as exc:
        raise FichierIllisible(f"fichier .pdf illisible ({type(exc).__name__})") from exc
    texte = "\n".join(pages)
    if not texte.strip():
        raise PdfSansTexte(
            "PDF sans texte extractible (document scanné probable) : refusé. "
            "Le sas ne fait pas d'OCR en v1 ; pseudonymiser l'image en la "
            "laissant lisible serait un faux sentiment de sécurité."
        )
    return texte


def extraire_texte(nom: str, octets: bytes) -> str:
    """Texte brut d'un fichier déposé, ou une ErreurFichier explicite."""
    extension = PurePosixPath(nom.lower()).suffix
    if extension not in EXTENSIONS_SUPPORTEES:
        raise FormatNonSupporte(
            f"format {extension or 'sans extension'!r} non supporté ; formats "
            f"acceptés : {', '.join(sorted(EXTENSIONS_SUPPORTEES))}. Un PDF "
            "scanné est refusé aussi : pas d'OCR en v1."
        )
    if extension in EXTENSIONS_TEXTE:
        return _texte_brut(octets)
    if extension == ".docx":
        return _texte_docx(octets)
    return _texte_pdf(octets)


# Propriétés textuelles recopiées du document source vers l'export
# (arbitrage Lot 15 : elles appartiennent à l'utilisateur), mais passées
# ELLES AUSSI à la détection : un nom ou un courriel en propriété ne doit
# pas fuir par la bande.
_PROPRIETES_TEXTE = (
    "author",
    "last_modified_by",
    "title",
    "subject",
    "comments",
    "keywords",
    "category",
)


def pseudonymiser_docx(
    octets: bytes, pseudonymiser: Callable
) -> tuple[bytes, dict[str, int]]:
    """Reconstruit le .docx avec les textes pseudonymisés (Lot 15).

    Paragraphes, cellules de tableaux et propriétés textuelles passent par
    `pseudonymiser` (déjà lié au dossier : mêmes placeholders que
    l'analyse). Un paragraphe sans détection garde sa mise en forme ; un
    paragraphe modifié la perd (limite documentée : python-docx remplace
    ses runs).
    """
    document = _ouvrir_docx(octets)
    comptes: Counter = Counter()

    def remplacer(texte: str) -> str:
        resultat = pseudonymiser(texte)
        comptes.update(resultat.comptes_par_type)
        return resultat.texte

    paragraphes = list(document.paragraphs)
    for table in document.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                paragraphes.extend(cellule.paragraphs)
    for paragraphe in paragraphes:
        if not paragraphe.text.strip():
            continue
        remplace = remplacer(paragraphe.text)
        if remplace != paragraphe.text:
            paragraphe.text = remplace

    proprietes = document.core_properties
    for champ in _PROPRIETES_TEXTE:
        valeur = getattr(proprietes, champ)
        if valeur:
            setattr(proprietes, champ, remplacer(valeur))

    tampon = io.BytesIO()
    document.save(tampon)
    return tampon.getvalue(), dict(comptes)
