"""Lot 15 : extraction de texte des fichiers déposés.

.txt / .md / .csv en texte brut, .docx (paragraphes ET tableaux), .pdf
textuel. Un PDF sans texte extractible (scanné) est refusé explicitement :
pas d'OCR en v1, un refus documenté vaut mieux qu'un texte silencieusement
vide qui laisserait croire à une pseudonymisation réussie. Tout contenu de
test est synthétique (REQ-009).
"""

import io

import pytest

from sas_confiance_ia.fichiers import (
    EXTENSIONS_SUPPORTEES,
    FichierIllisible,
    FormatNonSupporte,
    PdfSansTexte,
    extraire_texte,
)

TEXTE_FICTIF = "Conseil medical fictif. Agent : Karim Benhaddou, matricule 2014-0883."


def docx_fictif() -> bytes:
    """Un .docx synthétique : deux paragraphes et un tableau 2x2."""
    from docx import Document

    document = Document()
    document.add_paragraph("Conseil medical fictif du 11 juin 2026.")
    document.add_paragraph("Agent : Karim Benhaddou.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "NIR"
    table.cell(0, 1).text = "1 79 03 26 121 044 87"
    table.cell(1, 0).text = "Courriel"
    table.cell(1, 1).text = "k.benhaddou@exemple-fictif.fr"
    tampon = io.BytesIO()
    document.save(tampon)
    return tampon.getvalue()


def pdf_fictif(texte: str = TEXTE_FICTIF) -> bytes:
    """Un PDF textuel minimal assemblé à la main (offsets xref exacts)."""
    contenu = f"BT /F1 12 Tf 72 720 Td ({texte}) Tj ET".encode("latin-1")
    objets = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(contenu)).encode() + b" >>\nstream\n" + contenu + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    corps = io.BytesIO()
    corps.write(b"%PDF-1.4\n")
    positions = []
    for i, objet in enumerate(objets, start=1):
        positions.append(corps.tell())
        corps.write(f"{i} 0 obj\n".encode())
        corps.write(objet)
        corps.write(b"\nendobj\n")
    debut_xref = corps.tell()
    corps.write(f"xref\n0 {len(objets) + 1}\n".encode())
    corps.write(b"0000000000 65535 f \n")
    for position in positions:
        corps.write(f"{position:010d} 00000 n \n".encode())
    corps.write(
        f"trailer\n<< /Size {len(objets) + 1} /Root 1 0 R >>\n"
        f"startxref\n{debut_xref}\n%%EOF\n".encode()
    )
    return corps.getvalue()


def pdf_sans_texte() -> bytes:
    """Un PDF valide sans aucun opérateur de texte : un scan typique."""
    return pdf_fictif(texte="").replace(b"BT /F1 12 Tf 72 720 Td () Tj ET", b" " * 31)


# --- Texte brut --------------------------------------------------------------


def test_txt_md_csv_sont_lus_en_texte_brut():
    for nom in ["note.txt", "note.md", "tableau.csv"]:
        assert extraire_texte(nom, b"Karim Benhaddou, avenue de l'Europe.") == (
            "Karim Benhaddou, avenue de l'Europe."
        )


def test_le_texte_latin1_est_lu_en_repli():
    octets = "Décision notifiée à M. Benhaddou.".encode("latin-1")
    assert "Décision notifiée" in extraire_texte("note.txt", octets)


# --- DOCX --------------------------------------------------------------------


def test_docx_paragraphes_et_tableaux_sont_extraits():
    texte = extraire_texte("pv.docx", docx_fictif())
    assert "Conseil medical fictif du 11 juin 2026." in texte
    assert "Agent : Karim Benhaddou." in texte
    # Les tableaux ne sont jamais perdus : c'est là que vivent les NIR.
    assert "1 79 03 26 121 044 87" in texte
    assert "k.benhaddou@exemple-fictif.fr" in texte


def test_docx_corrompu_est_refuse_explicitement():
    with pytest.raises(FichierIllisible, match="docx"):
        extraire_texte("pv.docx", b"PK\x03\x04 pas un vrai docx")


# --- PDF ---------------------------------------------------------------------


def test_pdf_textuel_est_extrait():
    assert "Karim Benhaddou" in extraire_texte("pv.pdf", pdf_fictif())


def test_pdf_sans_texte_est_refuse_pas_d_ocr():
    # Refus documenté : un scan pseudonymisé « avec succès » serait un
    # mensonge dangereux (le texte visible de l'image partirait intact).
    with pytest.raises(PdfSansTexte, match="scann"):
        extraire_texte("scan.pdf", pdf_sans_texte())


def test_pdf_corrompu_est_refuse_explicitement():
    with pytest.raises(FichierIllisible, match="pdf"):
        extraire_texte("pv.pdf", b"%PDF-1.4 tronque")


# --- Garde-fous --------------------------------------------------------------


def test_format_inconnu_refuse_avec_la_liste_des_formats():
    with pytest.raises(FormatNonSupporte) as erreur:
        extraire_texte("tableur.xlsx", b"PK\x03\x04")
    assert ".docx" in str(erreur.value)


def test_les_extensions_supportees_sont_publiees():
    assert EXTENSIONS_SUPPORTEES == {".txt", ".md", ".csv", ".docx", ".pdf"}


def test_la_casse_de_l_extension_est_toleree():
    assert "Karim" in extraire_texte("PV.DOCX", docx_fictif())
