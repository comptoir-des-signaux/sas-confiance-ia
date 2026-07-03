"""Lot 15 : dépôt de fichiers dans l'UI et export .docx pseudonymisé.

Q3 (arbitrage) : le texte extrait du fichier de l'utilisateur peut revenir
à son navigateur (c'est SON document) ; la réponse serveur, elle, ne porte
jamais les valeurs détectées en mode sérieux : positions seulement. Le nom
du fichier n'entre jamais au journal (REQ-003 : il peut contenir un nom).
"""

import io
import logging

import pytest
from fastapi.testclient import TestClient

from sas_confiance_ia.api import creer_application
from sas_confiance_ia.backends import BackendCapture
from sas_confiance_ia.pseudonymiseur import Pseudonymiseur
from sas_confiance_ia.vault import VaultMemoire

from .test_fichiers import docx_fictif, pdf_sans_texte

NIR_FICTIF = "2 92 07 82 045 033 55"
EMAIL_FICTIF = "k.benhaddou@exemple-fictif.fr"


@pytest.fixture
def client():
    application = creer_application(
        pseudonymiseur=Pseudonymiseur(VaultMemoire()),
        backend=BackendCapture(),
        modeles=["modele-de-test"],
    )
    return TestClient(application)


def deposer(client, nom: str, octets: bytes, dossier: str = "d-fichier", mode: str = "serieux"):
    return client.post(
        "/ui/fichier",
        files={"fichier": (nom, io.BytesIO(octets))},
        data={"dossier_id": dossier, "mode": mode},
    )


# --- Dépôt et pseudonymisation -----------------------------------------------


def test_un_txt_depose_est_pseudonymise_avec_texte_origine(client):
    corps = deposer(client, "note.txt", f"Contacter {EMAIL_FICTIF}.".encode()).json()
    assert corps["texte_origine"] == f"Contacter {EMAIL_FICTIF}."
    assert corps["texte"] == "Contacter [EMAIL_001]."
    assert corps["nom"] == "note.txt"
    assert corps["comptes_par_type"] == {"EMAIL": 1}


def test_le_docx_est_pseudonymise_tableaux_compris(client):
    corps = deposer(client, "pv.docx", docx_fictif()).json()
    assert NIR_FICTIF in corps["texte_origine"]
    assert NIR_FICTIF not in corps["texte"]
    assert "[NIR_001]" in corps["texte"]
    assert "[EMAIL_001]" in corps["texte"]


def test_en_mode_serieux_les_detections_sont_des_positions_seulement(client):
    corps = deposer(client, "pv.docx", docx_fictif()).json()
    for detection in corps["detections"]:
        assert set(detection) == {"type", "debut", "fin", "score"}
    # Les positions se lisent dans le texte d'origine que la page possède.
    (position_nir,) = [d for d in corps["detections"] if d["type"] == "FR_NIR"]
    assert corps["texte_origine"][position_nir["debut"] : position_nir["fin"]] == NIR_FICTIF


def test_en_mode_demo_les_valeurs_sont_visibles(client):
    corps = deposer(client, "pv.docx", docx_fictif(), mode="demo").json()
    (detection_nir,) = [d for d in corps["detections"] if d["type"] == "FR_NIR"]
    assert detection_nir["valeur"] == NIR_FICTIF
    assert detection_nir["placeholder"] == "[NIR_001]"


def test_le_pdf_scanne_est_refuse_en_422_avec_explication(client):
    reponse = deposer(client, "scan.pdf", pdf_sans_texte())
    assert reponse.status_code == 422
    assert "OCR" in reponse.json()["detail"]


def test_le_format_inconnu_est_refuse_en_415(client):
    reponse = deposer(client, "tableur.xlsx", b"PK\x03\x04")
    assert reponse.status_code == 415
    assert ".docx" in reponse.json()["detail"]


def test_le_nom_du_fichier_n_entre_jamais_au_journal(client, caplog):
    # REQ-003 : « dossier_karim_benhaddou.docx » est une donnée personnelle.
    with caplog.at_level(logging.INFO, logger="sas_confiance_ia.journal"):
        deposer(client, "dossier_karim_benhaddou.docx", docx_fictif())
    assert caplog.messages, "le dépôt doit être journalisé"
    assert all("dossier_karim_benhaddou" not in message for message in caplog.messages)


def test_un_dossier_demo_est_refuse_sur_le_chemin_serieux(client):
    deposer(client, "note.txt", b"Texte de demonstration.", mode="demo")
    reponse = deposer(client, "note.txt", b"Suite.", mode="serieux")
    assert reponse.status_code == 409


# --- Export .docx pseudonymisé -----------------------------------------------


def exporter(client, nom: str, octets: bytes, dossier: str = "d-fichier"):
    return client.post(
        "/ui/fichier/export-docx",
        files={"fichier": (nom, io.BytesIO(octets))},
        data={"dossier_id": dossier},
    )


def test_l_export_docx_ne_contient_plus_aucune_valeur_detectee(client):
    from docx import Document

    reponse = exporter(client, "pv.docx", docx_fictif())
    assert reponse.status_code == 200
    document = Document(io.BytesIO(reponse.content))
    textes = [p.text for p in document.paragraphs] + [
        cellule.text for table in document.tables for ligne in table.rows for cellule in ligne.cells
    ]
    tout = "\n".join(textes)
    assert NIR_FICTIF not in tout
    assert EMAIL_FICTIF not in tout
    assert "[NIR_001]" in tout
    assert "[EMAIL_001]" in tout


def test_l_export_et_l_analyse_partagent_les_placeholders_du_dossier(client):
    # Même dossier, même vault : l'export reprend les placeholders déjà émis.
    from docx import Document

    corps = deposer(client, "pv.docx", docx_fictif(), dossier="d-coherent").json()
    assert "[NIR_001]" in corps["texte"]
    reponse = exporter(client, "pv.docx", docx_fictif(), dossier="d-coherent")
    document = Document(io.BytesIO(reponse.content))
    tout = "\n".join(
        cellule.text for table in document.tables for ligne in table.rows for cellule in ligne.cells
    )
    assert "[NIR_001]" in tout
    assert "[NIR_002]" not in tout


def test_les_metadonnees_source_sont_recopiees_et_pseudonymisees(client):
    # Arbitrage lot 15 : l'export hérite des propriétés du document source
    # (elles appartiennent à l'utilisateur), mais elles passent AUSSI par la
    # détection : un courriel en commentaire ne doit pas fuir.
    import io as io_

    from docx import Document

    source = Document(io_.BytesIO(docx_fictif()))
    source.core_properties.author = "Service RH fictif"
    source.core_properties.title = "PV de conseil medical fictif"
    source.core_properties.comments = f"Relu par {EMAIL_FICTIF}"
    tampon = io_.BytesIO()
    source.save(tampon)

    reponse = exporter(client, "pv.docx", tampon.getvalue())
    exporte = Document(io_.BytesIO(reponse.content))
    assert exporte.core_properties.author == "Service RH fictif"
    assert exporte.core_properties.title == "PV de conseil medical fictif"
    assert EMAIL_FICTIF not in exporte.core_properties.comments
    assert "[EMAIL" in exporte.core_properties.comments


def test_l_export_reste_reversible_via_la_reidentification(client):
    from docx import Document

    reponse = exporter(client, "pv.docx", docx_fictif(), dossier="d-retour")
    document = Document(io.BytesIO(reponse.content))
    pseudonymise = document.paragraphs[1].text
    assert pseudonymise == "Agent : Karim Benhaddou."  # PERSONNE exige le NER : inchangé ici
    cellule = document.tables[0].cell(0, 1).text
    corps = client.post(
        "/ui/reidentifier",
        json={"texte": cellule, "dossier_id": "d-retour"},
    ).json()
    assert corps["texte"] == NIR_FICTIF


def test_l_export_refuse_tout_sauf_docx(client):
    reponse = exporter(client, "note.txt", b"du texte")
    assert reponse.status_code == 415
    assert "docx" in reponse.json()["detail"]


# --- Page Fichiers -------------------------------------------------------------


def test_la_page_fichiers_est_servie(client):
    reponse = client.get("/fichiers")
    assert reponse.status_code == 200
    assert "text/html" in reponse.headers["content-type"]
    page = reponse.text
    assert "déposer" in page.lower() or "glisser" in page.lower()
    # Côte à côte origine / pseudonymisé, surlignage côté client (Q3).
    assert "cote-a-cote" in page
    assert "surligner" in page


def test_l_accueil_et_la_page_fichiers_se_referencent(client):
    assert '"/fichiers"' in client.get("/").text
    assert 'href="/"' in client.get("/fichiers").text
